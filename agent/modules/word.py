"""
Word automation module — read, create, edit, and extract from .docx files.

Contract: exposes TOOL_DEFINITIONS and dispatch() for the registry.
"""

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


TOOL_DEFINITIONS: list[dict] = [
    {
        "name": "word_read",
        "description": "Read the full text content of a .docx file, paragraph by paragraph.",
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Absolute or relative path to the .docx file."},
                "include_headings": {
                    "type": "boolean",
                    "description": "If true, prefix each heading paragraph with its level (H1, H2, …). Default true.",
                    "default": True,
                },
            },
            "required": ["path"],
        },
    },
    {
        "name": "word_create",
        "description": (
            "Create a new .docx file from a list of content blocks. "
            "Each block is a dict with 'type' (heading|paragraph|bullet|table) and 'text' or 'rows'."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Output file path (will be created/overwritten)."},
                "title": {"type": "string", "description": "Document title (H1 heading at the top)."},
                "blocks": {
                    "type": "array",
                    "description": "List of content blocks.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "type": {
                                "type": "string",
                                "enum": ["heading", "paragraph", "bullet", "table"],
                                "description": "'heading' uses level field (1-9), 'table' uses rows field (list of list of str).",
                            },
                            "text": {"type": "string"},
                            "level": {"type": "integer", "description": "Heading level 1-9 (default 2)."},
                            "bold": {"type": "boolean"},
                            "rows": {
                                "type": "array",
                                "items": {"type": "array", "items": {"type": "string"}},
                                "description": "For type=table: list of rows, each row is a list of cell strings.",
                            },
                        },
                        "required": ["type"],
                    },
                },
            },
            "required": ["path", "blocks"],
        },
    },
    {
        "name": "word_append",
        "description": "Append paragraphs or a table to an existing .docx file.",
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Path to existing .docx file."},
                "blocks": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "type": {"type": "string", "enum": ["heading", "paragraph", "bullet", "table"]},
                            "text": {"type": "string"},
                            "level": {"type": "integer"},
                            "bold": {"type": "boolean"},
                            "rows": {"type": "array", "items": {"type": "array", "items": {"type": "string"}}},
                        },
                        "required": ["type"],
                    },
                },
            },
            "required": ["path", "blocks"],
        },
    },
    {
        "name": "word_find_replace",
        "description": "Find and replace all occurrences of a text string in a .docx file.",
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {"type": "string"},
                "find": {"type": "string", "description": "Text to find."},
                "replace": {"type": "string", "description": "Replacement text."},
                "case_sensitive": {"type": "boolean", "default": True},
            },
            "required": ["path", "find", "replace"],
        },
    },
    {
        "name": "word_get_metadata",
        "description": "Get metadata (author, title, subject, word count, paragraph count) from a .docx file.",
        "input_schema": {
            "type": "object",
            "properties": {
                "path": {"type": "string"},
            },
            "required": ["path"],
        },
    },
]


# ---------------------------------------------------------------------------
# Implementations
# ---------------------------------------------------------------------------

def _add_block(doc: Document, block: dict) -> None:
    btype = block.get("type", "paragraph")
    text  = block.get("text", "")
    bold  = block.get("bold", False)

    if btype == "heading":
        level = block.get("level", 2)
        doc.add_heading(text, level=level)

    elif btype == "paragraph":
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True

    elif btype == "bullet":
        doc.add_paragraph(text, style="List Bullet")

    elif btype == "table":
        rows = block.get("rows", [])
        if not rows:
            return
        cols = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=cols)
        tbl.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                tbl.cell(r_idx, c_idx).text = str(cell_text)


def word_read(path: str, include_headings: bool = True) -> dict[str, Any]:
    doc = Document(path)
    paragraphs = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        style = p.style.name
        if include_headings and style.startswith("Heading"):
            level = style.split()[-1] if style.split()[-1].isdigit() else "?"
            paragraphs.append({"type": f"H{level}", "text": p.text})
        else:
            paragraphs.append({"type": "paragraph", "text": p.text})
    return {
        "path": path,
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
    }


def word_create(path: str, blocks: list[dict], title: str = "") -> dict[str, Any]:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for block in blocks:
        _add_block(doc, block)
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return {"path": path, "status": "created", "blocks_written": len(blocks)}


def word_append(path: str, blocks: list[dict]) -> dict[str, Any]:
    doc = Document(path)
    for block in blocks:
        _add_block(doc, block)
    doc.save(path)
    return {"path": path, "status": "appended", "blocks_written": len(blocks)}


def word_find_replace(path: str, find: str, replace: str, case_sensitive: bool = True) -> dict[str, Any]:
    doc = Document(path)
    count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            original = run.text
            if case_sensitive:
                if find in run.text:
                    run.text = run.text.replace(find, replace)
                    count += run.text.count(replace)
            else:
                import re
                new_text, n = re.subn(re.escape(find), replace, run.text, flags=re.IGNORECASE)
                run.text = new_text
                count += n
    doc.save(path)
    return {"path": path, "replacements": count, "status": "done"}


def word_get_metadata(path: str) -> dict[str, Any]:
    doc = Document(path)
    core = doc.core_properties
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    return {
        "path": path,
        "title": core.title or "",
        "author": core.author or "",
        "subject": core.subject or "",
        "created": str(core.created) if core.created else "",
        "modified": str(core.modified) if core.modified else "",
        "paragraph_count": len(doc.paragraphs),
        "word_count": word_count,
    }


def dispatch(tool_name: str, tool_input: dict) -> dict[str, Any]:
    match tool_name:
        case "word_read":         return word_read(**tool_input)
        case "word_create":       return word_create(**tool_input)
        case "word_append":       return word_append(**tool_input)
        case "word_find_replace": return word_find_replace(**tool_input)
        case "word_get_metadata": return word_get_metadata(**tool_input)
        case _: raise ValueError(f"Unknown tool: {tool_name}")
